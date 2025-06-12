from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from htmldocx import HtmlToDocx
import logging
import re
from typing import List, Dict, Any
from bs4 import BeautifulSoup
from app.utils.logging_config import pdf_vision_logger, generate_request_id

logger = logging.getLogger(__name__)

def parse_html_to_word_format(html_content: str) -> List[Dict[str, Any]]:
    """
    Convert HTML content to Word document format using htmldocx library with direct alignment extraction.
    This handles all HTML formatting including CKEditor output with inline styles.
    """
    if not html_content or html_content.strip() == "":
        logger.debug("Empty HTML content provided")
        return []
    
    try:
        logger.debug(f"Converting HTML to Word format: {html_content[:200]}...")
        
        # Parse HTML with BeautifulSoup to extract structured information
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Find all text elements with their formatting
        text_elements = []
        
        # First, handle div containers with alignment that contain paragraphs
        for div_element in soup.find_all('div'):
            div_alignment = "left"
            
            # Check if div has alignment styles
            if div_element.get('style'):
                style = div_element.get('style', '')
                if 'text-align: center' in style or 'text-align:center' in style:
                    div_alignment = "center"
                elif 'text-align: right' in style or 'text-align:right' in style:
                    div_alignment = "right"
                elif 'text-align: justify' in style or 'text-align:justify' in style:
                    div_alignment = "justify"
            
            # If div has alignment, apply it to child paragraphs
            if div_alignment != "left":
                child_paragraphs = div_element.find_all('p', recursive=False)
                for p in child_paragraphs:
                    text = p.get_text(strip=True)
                    if not text:
                        continue
                    
                    # Inherit alignment from parent div unless paragraph has its own
                    p_alignment = div_alignment
                    if p.get('style'):
                        p_style = p.get('style', '')
                        if 'text-align: center' in p_style or 'text-align:center' in p_style:
                            p_alignment = "center"
                        elif 'text-align: right' in p_style or 'text-align:right' in p_style:
                            p_alignment = "right"
                        elif 'text-align: justify' in p_style or 'text-align:justify' in p_style:
                            p_alignment = "justify"
                    
                    # Extract formatting
                    bold = bool(p.find(['strong', 'b']) or (p.get('style') and ('font-weight: bold' in p.get('style') or 'font-weight:bold' in p.get('style'))))
                    italic = bool(p.find(['em', 'i']) or (p.get('style') and ('font-style: italic' in p.get('style') or 'font-style:italic' in p.get('style'))))
                    
                    text_elements.append({
                        "text": text,
                        "alignment": p_alignment,
                        "bold": bold,
                        "italic": italic,
                        "font_size": 12,
                        "element_type": "p"
                    })
                    
                    logger.debug(f"Extracted from div: '{text[:30]}...' -> alignment={p_alignment}, bold={bold}, italic={italic}")
        
        # Then handle standalone paragraphs and other elements that aren't inside divs with alignment
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            # Skip paragraphs that are already processed as children of aligned divs
            parent_div = element.find_parent('div')
            if parent_div and parent_div.get('style') and 'text-align:' in parent_div.get('style'):
                continue  # Already processed above
            
            text = element.get_text(strip=True)
            if not text:
                continue
                
            alignment = "left"  # default
            
            # Check for inline styles (style="text-align: center;")
            if element.get('style'):
                style = element.get('style', '')
                if 'text-align: center' in style or 'text-align:center' in style:
                    alignment = "center"
                elif 'text-align: right' in style or 'text-align:right' in style:
                    alignment = "right"
                elif 'text-align: justify' in style or 'text-align:justify' in style:
                    alignment = "justify"
            
            # Check for CSS classes (class="ql-align-center")
            if element.get('class'):
                classes = element.get('class', [])
                if isinstance(classes, list):
                    if 'ql-align-center' in classes:
                        alignment = "center"
                    elif 'ql-align-right' in classes:
                        alignment = "right"
                    elif 'ql-align-justify' in classes:
                        alignment = "justify"
                elif isinstance(classes, str):
                    if 'ql-align-center' in classes:
                        alignment = "center"
                    elif 'ql-align-right' in classes:
                        alignment = "right"
                    elif 'ql-align-justify' in classes:
                        alignment = "justify"
            
            # Extract formatting information directly from element
            bold = False
            italic = False
            
            # Check for strong/b tags
            if element.find(['strong', 'b']) or (element.get('style') and ('font-weight: bold' in element.get('style') or 'font-weight:bold' in element.get('style'))):
                bold = True
            
            # Check for em/i tags
            if element.find(['em', 'i']) or (element.get('style') and ('font-style: italic' in element.get('style') or 'font-style:italic' in element.get('style'))):
                italic = True
            
            text_elements.append({
                "text": text,
                "alignment": alignment,
                "bold": bold,
                "italic": italic,
                "font_size": 12,
                "element_type": element.name
            })
            
            logger.debug(f"Extracted: '{text[:30]}...' -> alignment={alignment}, bold={bold}, italic={italic}")
        
        # If BeautifulSoup extraction found elements, use them directly
        if text_elements:
            logger.debug(f"Using direct BeautifulSoup extraction: {len(text_elements)} text blocks")
            return text_elements
        
        # Fallback: Use htmldocx as before but with improved text matching
        alignment_map = {}
        elements = soup.find_all(True)
        
        for i, element in enumerate(elements):
            text = element.get_text(strip=True)
            if not text:
                continue
                
            alignment = "left"  # default
            
            # Check for inline styles (style="text-align: center;")
            if element.get('style'):
                style = element.get('style', '')
                if 'text-align: center' in style or 'text-align:center' in style:
                    alignment = "center"
                elif 'text-align: right' in style or 'text-align:right' in style:
                    alignment = "right"
                elif 'text-align: justify' in style or 'text-align:justify' in style:
                    alignment = "justify"
            
            # Check for QuillTextEditor CSS classes (class="ql-align-center")
            if element.get('class'):
                classes = element.get('class', [])
                if isinstance(classes, list):
                    if 'ql-align-center' in classes:
                        alignment = "center"
                    elif 'ql-align-right' in classes:
                        alignment = "right"
                    elif 'ql-align-justify' in classes:
                        alignment = "justify"
                elif isinstance(classes, str):
                    if 'ql-align-center' in classes:
                        alignment = "center"
                    elif 'ql-align-right' in classes:
                        alignment = "right"
                    elif 'ql-align-justify' in classes:
                        alignment = "justify"
            
            if alignment != "left":
                # Store alignment for multiple possible text variations
                alignment_map[text] = alignment
                alignment_map[text.strip()] = alignment
                alignment_map[text.replace('\n', ' ')] = alignment
                alignment_map[text.replace('\n', ' ').strip()] = alignment
                logger.debug(f"Found {alignment} alignment for: '{text[:50]}...' (element: {element.name})")
        
        # Create a new document and use htmldocx
        document = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html_content, document)
        
        # Convert document back to our block format
        blocks = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text = paragraph.text.strip()
                
                # Check if this text has alignment in our map (try multiple variations)
                alignment = (alignment_map.get(text) or 
                           alignment_map.get(text.strip()) or 
                           alignment_map.get(text.replace('\n', ' ')) or
                           alignment_map.get(text.replace('\n', ' ').strip()) or
                           "left")
                
                # Check for formatting in runs
                bold = any(run.bold for run in paragraph.runs if run.bold)
                italic = any(run.italic for run in paragraph.runs if run.italic)
                
                blocks.append({
                    "text": text,
                    "alignment": alignment,
                    "bold": bold,
                    "italic": italic,
                    "font_size": 12  # Default size, htmldocx handles actual sizing
                })
                
                logger.debug(f"Block: '{text[:30]}...' -> alignment={alignment}, bold={bold}")
        
        logger.debug(f"Successfully converted HTML to {len(blocks)} text blocks")
        return blocks
        
    except Exception as e:
        logger.error(f"Error converting HTML with htmldocx: {str(e)}")
        # Fallback to simple text extraction
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()
            if text.strip():
                return [{
                    "text": text.strip(),
                    "alignment": "left",
                    "bold": False,
                    "italic": False,
                    "font_size": 12
                }]
        except Exception as fallback_error:
            logger.error(f"Fallback parsing also failed: {str(fallback_error)}")
        
        return []

class WordGenerator:
    def __init__(self, formatted_text: List[Dict]):
        self.formatted_text = formatted_text
        self.logger = logging.getLogger(__name__)
        
    def generate_document(self, output_path: str):
        """
        Create a Word document from formatted text stored in self.formatted_text.
        
        Args:
            output_path (str): Path where the Word document will be saved
        """
        try:
            self.logger.info(f"Creating Word document at: {output_path}")
            self.logger.info(f"Input contains {len(self.formatted_text)} text elements")
            
            page_counts = {}
            for item in self.formatted_text:
                page = item.get("page", "unknown")
                if page not in page_counts:
                    page_counts[page] = 0
                page_counts[page] += 1
            
            self.logger.info(f"Text elements by page: {page_counts}")
            
            validated_text = []
            for item in self.formatted_text:
                # Skip items without text
                if not item.get("text"):
                    continue
                    
                # Ensure all required keys exist
                clean_item = {
                    "text": item.get("text", ""),
                    "font": item.get("font", "Calibri"),
                    "size": item.get("size", 11),
                    # Default to black color for text visibility
                    "color": (0, 0, 0),  
                    "is_bold": item.get("is_bold", False),
                    "is_italic": item.get("is_italic", False),
                    "alignment": item.get("alignment", "left"),  # ✅ CRITICAL FIX: Preserve alignment!
                    "page": item.get("page", 1),
                    "block_no": item.get("block_no", 0),
                    "line_no": item.get("line_no", 0),
                    "is_last_span_in_line": item.get("is_last_span_in_line", False),
                    "bbox": item.get("bbox", [0, 0, 100, 20])
                }
                
                # Only add the original color if it's from the original PDF extraction
                # This helps avoid white text from edited content
                if "color" in item and isinstance(item["color"], (list, tuple)) and len(item["color"]) >= 3:
                    # Check if the color is too light (close to white)
                    r, g, b = item["color"][:3]
                    
                    # Convert to values between 0-255 if they're not already
                    if isinstance(r, float) and r <= 1.0:
                        r = int(r * 255)
                    if isinstance(g, float) and g <= 1.0:
                        g = int(g * 255)
                    if isinstance(b, float) and b <= 1.0:
                        b = int(b * 255)
                    
                    # Calculate brightness (higher values are lighter)
                    brightness = (r + g + b) / 3
                    
                    # If the color is too light (close to white), use black instead
                    if brightness < 230:  # Threshold for "not too light"
                        clean_item["color"] = item["color"]
                
                validated_text.append(clean_item)
                
            if not validated_text:
                self.logger.warning("No valid text to export to Word document")
                # Create empty document with a message
                doc = Document()
                doc.add_paragraph("No text content could be extracted from the PDF.")
                doc.save(output_path)
                self.logger.info("Created empty document with message")
                return
                
            # Continue with normal processing using validated text
            self.logger.info(f"Processing {len(validated_text)} validated text elements")
            doc = Document()
            
            # Group text elements by page and block for better structure
            text_by_blocks = self._group_by_page_and_block(validated_text)
            self.logger.info(f"Grouped text into {len(text_by_blocks)} pages")
            
            # Process each page and block
            for page_num in sorted(text_by_blocks.keys()):
                self.logger.info(f"Processing page {page_num} with {len(text_by_blocks[page_num])} blocks")
                
                # Check if this page has any content before adding page break
                page_has_content = False
                for block_no in text_by_blocks[page_num]:
                    spans = text_by_blocks[page_num][block_no]
                    for span in spans:
                        if span.get("text", "").strip():
                            page_has_content = True
                            break
                    if page_has_content:
                        break
                
                # Only add page break if this page has content and it's not the first page
                if page_num > 1 and page_has_content:
                    doc.add_page_break()
                
                # Skip this page entirely if it has no content
                if not page_has_content:
                    self.logger.info(f"Skipping page {page_num} - no content")
                    continue
                
                # Process blocks on this page
                for block_no in sorted(text_by_blocks[page_num].keys()):
                    # Create a new paragraph for this block
                    current_paragraph = doc.add_paragraph()
                    current_line_no = -1
                    
                    # Get spans for this block
                    spans = text_by_blocks[page_num][block_no]
                    
                    # Sort spans by line and position
                    spans.sort(key=lambda x: (x.get("line_no", 0), x.get("bbox", [0, 0, 0, 0])[0]))
                    
                    # ALWAYS check for alignment first, regardless of number of spans
                    # Get alignment from the first span (all spans in a block should have same alignment)
                    first_span = spans[0] if spans else {}
                    alignment = first_span.get("alignment")
                    
                    # DEBUG: Force print to console (bypassing logger issues)
                    print(f"🔍 DEBUG WordGenerator: Block '{first_span.get('text', '')[:30]}...' has alignment='{alignment}'")
                    
                    # Apply paragraph-level alignment to ALL blocks
                    if alignment == "center":
                        current_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        print(f"✅ Applied CENTER alignment to block: '{first_span.get('text', '')[:30]}...'")
                        self.logger.info(f"Applied CENTER alignment to block: '{first_span.get('text', '')[:30]}...'")
                    elif alignment == "right":
                        current_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        print(f"✅ Applied RIGHT alignment to block: '{first_span.get('text', '')[:30]}...'")
                        self.logger.info(f"Applied RIGHT alignment to block: '{first_span.get('text', '')[:30]}...'")
                    elif alignment == "justify":
                        current_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        print(f"✅ Applied JUSTIFY alignment to block: '{first_span.get('text', '')[:30]}...'")
                        self.logger.info(f"Applied JUSTIFY alignment to block: '{first_span.get('text', '')[:30]}...'")
                    elif first_span.get("is_indent"):
                        # Add indentation for indent blocks
                        from docx.shared import Inches
                        current_paragraph.paragraph_format.left_indent = Inches(0.5)
                        print(f"✅ Applied INDENT to block: '{first_span.get('text', '')[:30]}...'")
                        self.logger.info(f"Applied INDENT to block: '{first_span.get('text', '')[:30]}...'")
                    else:
                        print(f"❌ No special alignment for block: '{first_span.get('text', '')[:30]}...' (alignment={alignment})")
                        self.logger.info(f"No special alignment for block: '{first_span.get('text', '')[:30]}...' (alignment={alignment})")
                    
                    # Now process the spans - simplified logic since alignment is handled above
                    if len(spans) == 1:
                        span = spans[0]
                        
                        # Debug: Show what formatting we're applying
                        self.logger.info(f"Single span block: '{span.get('text', '')[:30]}...' -> bold={span.get('is_bold')}, size={span.get('size')}")
                        
                        # Add the text with appropriate formatting
                        if span.get("text"):
                            run = current_paragraph.add_run(span.get("text", ""))
                            
                            # Apply text formatting
                            font = run.font
                            font.name = span.get("font", "Calibri")
                            font.size = Pt(span.get("size", 11))
                            font.bold = span.get("is_bold", False)
                            font.italic = span.get("is_italic", False)
                            font.color.rgb = RGBColor(0, 0, 0)  # Black
                            
                            # Apply color if available
                            if "color" in span and span["color"]:
                                self._apply_color_to_run(font, span["color"])
                    else:
                        # Process multiple spans in the block (regular paragraph)
                        for span in spans:
                            # Skip empty spans
                            if not span.get("text"):
                                continue
                            
                            # Check if we need to add a line break within the same paragraph
                            if "line_no" in span and span["line_no"] != current_line_no and current_line_no != -1:
                                # Add line break instead of paragraph break
                                current_paragraph.add_run().add_break()
                            
                            current_line_no = span.get("line_no", 0)
                            
                            # Create text run with proper formatting
                            run = current_paragraph.add_run(span.get("text", "").rstrip())
                            
                            # Apply formatting to this specific run only
                            font = run.font
                            font.name = span.get("font", "Calibri")
                            font.size = Pt(span.get("size", 11))
                            font.bold = span.get("is_bold", False)
                            font.italic = span.get("is_italic", False)
                            font.color.rgb = RGBColor(0, 0, 0)
                            
                            # Set color if available and not too light
                            if "color" in span and span["color"]:
                                self._apply_color_to_run(font, span["color"])
                            
                            # Add space after text if not the last span in a line and there's no space already
                            span_text = span.get("text", "")
                            if not span.get("is_last_span_in_line", False) and not span_text.endswith(" "):
                                current_paragraph.add_run(" ")
            
            # Save the document
            try:
                doc.save(output_path)
                self.logger.info(f"Word document saved to {output_path}")
            except Exception as save_error:
                self.logger.error(f"Error saving Word document: {str(save_error)}")
                self.logger.error(f"Exception type: {type(save_error).__name__}")
                self.logger.error(f"Stack trace: {traceback.format_exc()}")
                raise
            
        except Exception as e:
            self.logger.error(f"Error creating Word document: {str(e)}")
            self.logger.error(f"Exception type: {type(e).__name__}")
            self.logger.error(f"Stack trace: {traceback.format_exc()}")
            
            # Try to create a basic document with error information
            try:
                doc = Document()
                doc.add_heading("Error Report", 0)
                doc.add_paragraph(f"Error creating Word document: {str(e)}")
                doc.add_paragraph(f"Exception type: {type(e).__name__}")
                doc.add_heading("Stack Trace:", 1)
                doc.add_paragraph(traceback.format_exc())
                
                doc.add_heading("Input Data Summary:", 1)
                doc.add_paragraph(f"Total items: {len(self.formatted_text)}")
                
                # Sample of the first few items
                doc.add_heading("Sample Input Items:", 2)
                for i, item in enumerate(self.formatted_text[:5]):
                    doc.add_paragraph(f"Item {i}: {str(item)[:200]}...")
                    
                # Save error report
                error_path = output_path.replace(".docx", "_error_report.docx")
                doc.save(error_path)
                self.logger.info(f"Created error report at: {error_path}")
            except Exception as report_error:
                self.logger.error(f"Failed to create error report: {str(report_error)}")
            
            raise
            
    def add_metadata(self, doc: Document, title: str, author: str = None):
        """
        Add metadata to the Word document.
        
        Args:
            doc (Document): The Word document
            title (str): Document title
            author (str, optional): Document author
        """
        try:
            core_properties = doc.core_properties
            core_properties.title = title
            if author:
                core_properties.author = author
                
        except Exception as e:
            self.logger.error(f"Error adding metadata: {str(e)}")
            raise 

    def _group_by_page_and_block(self, formatted_text: List[Dict]) -> Dict:
        """
        Group text elements by page and block for easier processing.
        
        Args:
            formatted_text (List[Dict]): List of text elements with formatting
            
        Returns:
            Dict: Dictionary with structure {page_num: {block_no: [spans]}}
        """
        result = {}
        
        for span in formatted_text:
            page_num = span.get("page", 1)
            block_no = span.get("block_no", 0)
            
            # Initialize dictionaries if needed
            if page_num not in result:
                result[page_num] = {}
            if block_no not in result[page_num]:
                result[page_num][block_no] = []
            
            # Add span to the appropriate block
            result[page_num][block_no].append(span)
        
        return result

    def _apply_color_to_run(self, font, color_value):
        """
        Apply color to a text run, handling different color formats.
        
        Args:
            font: The font object to apply color to
            color_value: The color value (can be tuple, list, int)
        """
        try:
            # Handle different color formats
            if isinstance(color_value, (list, tuple)) and len(color_value) >= 3:
                # Try to convert RGB values and ensure they're within 0-255 range
                try:
                    # Get RGB values and scale to 0-255 if needed
                    r = int(color_value[0] * 255) if isinstance(color_value[0], float) and color_value[0] <= 1.0 else int(color_value[0])
                    g = int(color_value[1] * 255) if isinstance(color_value[1], float) and color_value[1] <= 1.0 else int(color_value[1])
                    b = int(color_value[2] * 255) if isinstance(color_value[2], float) and color_value[2] <= 1.0 else int(color_value[2])
                    
                    # Ensure values are in valid range
                    r = max(0, min(r, 255))
                    g = max(0, min(g, 255))
                    b = max(0, min(b, 255))
                    
                    # Calculate brightness to avoid very light colors
                    brightness = (r + g + b) / 3
                    
                    # If color is too light (close to white), use black instead
                    if brightness > 230:
                        self.logger.info(f"Converting light color {(r,g,b)} to black for better visibility")
                        font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        font.color.rgb = RGBColor(r, g, b)
                except (ValueError, TypeError):
                    # If conversion fails, use default black
                    self.logger.warning(f"Invalid color value: {color_value}, using black instead")
                    font.color.rgb = RGBColor(0, 0, 0)
            elif isinstance(color_value, (int, float)):
                # Grayscale value
                try:
                    gray = int(color_value * 255) if isinstance(color_value, float) and color_value <= 1.0 else int(color_value)
                    gray = max(0, min(gray, 255))  # Ensure in valid range
                    
                    # If grayscale value is too light, use black instead
                    if gray > 230:
                        self.logger.info(f"Converting light grayscale {gray} to black for better visibility")
                        font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        font.color.rgb = RGBColor(gray, gray, gray)
                except (ValueError, TypeError):
                    # If conversion fails, use default black
                    self.logger.warning(f"Invalid grayscale value: {color_value}, using black instead")
                    font.color.rgb = RGBColor(0, 0, 0)
            else:
                # For any other case, use default black
                font.color.rgb = RGBColor(0, 0, 0)
        except Exception as e:
            # Catch any other errors and use black as fallback
            self.logger.error(f"Error applying color: {str(e)}")
            font.color.rgb = RGBColor(0, 0, 0) 